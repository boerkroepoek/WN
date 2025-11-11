import openpyxl
from docx import Document
import os
# 'copy' is niet langer nodig
# from copy import deepcopy

def laad_vervangingen(excel_pad: str) -> dict | None:
    """
    Laadt termen en vervangingen uit het Excel-bestand.
    Start vanaf rij 2. Kolom A = Term (incl. {}), Kolom B = Input.
    """
    vervangingen = {}
    try:
        workbook = openpyxl.load_workbook(excel_pad, data_only=True)
        sheet = workbook.active
        
        for row in sheet.iter_rows(min_row=2, max_col=2, values_only=True):
            term = row[0]
            input_waarde = row[1]
            
            if term is not None:
                vervangingen[term] = str(input_waarde) if input_waarde is not None else ""
                
        print(f"✅ {len(vervangingen)} vervangingen geladen uit {excel_pad}")
        return vervangingen
        
    except FileNotFoundError:
        print(f"❌ FOUT: Excel-bestand niet gevonden op: {excel_pad}")
        return None
    except Exception as e:
        print(f"❌ FOUT bij het lezen van Excel: {e}")
        return None

# --- VERNIEUWDE, ROBUUSTE HELPER-FUNCTIE ---
def vervang_tekst_met_behoud_van_opmaak(para, term, input_waarde):
    """
    Vervangt tekst in een paragraaf, zelfs als deze is opgesplitst
    over meerdere 'runs', met behoud van de opmaak van de eerste run.
    
    Deze functie vervangt de term zo vaak als deze voorkomt.
    """
    # 1. Verzamel alle tekst en run-indices in een lijst
    # We hebben dit nodig om de 'full_text' indexen terug te mappen naar 'runs'
    runs_info = []
    full_text = ""
    for i, run in enumerate(para.runs):
        runs_info.append({'index': i, 'text': run.text})
        full_text += run.text

    # 2. Blijf vervangen zolang de term in de volledige tekst voorkomt
    while term in full_text:
        # 3. Vind de start- en eindindex van de term in de *volledige* tekst
        start_index = full_text.find(term)
        end_index = start_index + len(term)

        # 4. Map deze 'full_text' indexen terug naar de specifieke runs
        start_run_idx, start_char_idx = None, None
        end_run_idx, end_char_idx = None, None
        
        current_pos = 0
        # Loop door onze verzamelde run-info
        for run_info in runs_info:
            run_len = len(run_info['text'])
            
            # Check of de START-positie in deze run valt
            if start_run_idx is None and start_index < (current_pos + run_len):
                start_run_idx = run_info['index']
                start_char_idx = start_index - current_pos
            
            # Check of de EIND-positie in deze run valt
            if end_run_idx is None and end_index <= (current_pos + run_len):
                end_run_idx = run_info['index']
                end_char_idx = end_index - current_pos
                break # We hebben alles gevonden
                
            current_pos += run_len
            
        # Veiligheidscheck (zou niet moeten gebeuren als 'term in full_text')
        if start_run_idx is None or end_run_idx is None:
            print(f"⚠️ Waarschuwing: Kon '{term}' niet mappen naar runs. Stoppen.")
            break 

        # 5. Voer de vervanging uit (MET DE BUGFIX)
        
        # Haal de tekst op vóór de term (uit de start-run)
        tekst_voor = para.runs[start_run_idx].text[:start_char_idx]
        # Haal de tekst op ná de term (uit de eind-run)
        tekst_na = para.runs[end_run_idx].text[end_char_idx:]

        if start_run_idx == end_run_idx:
            # *** DE BUGFIX ***
            # Alles gebeurt in één run. Combineer alles in deze ene run.
            para.runs[start_run_idx].text = tekst_voor + input_waarde + tekst_na
        else:
            # De vervanging is verdeeld over meerdere runs
            
            # 1. Update de eerste run
            para.runs[start_run_idx].text = tekst_voor + input_waarde
            
            # 2. Maak alle tussenliggende runs leeg
            for i in range(start_run_idx + 1, end_run_idx):
                para.runs[i].text = ""
                
            # 3. Update de laatste run
            para.runs[end_run_idx].text = tekst_na
        
        # 6. Update de 'full_text' en 'runs_info' voor de *volgende* 'while'-loop
        # Dit is nodig zodat we de term opnieuw kunnen vinden als deze
        # meerdere keren voorkomt in de paragraaf.
        runs_info = []
        full_text = ""
        for i, run in enumerate(para.runs):
            runs_info.append({'index': i, 'text': run.text})
            full_text += run.text
            
        # Als de term niet meer in de (nieuwe) full_text zit, stopt de loop.


# --- AANGEPASTE VERWERK-FUNCTIE ---
def verwerk_paragrafen(paragrafen, vervangingen: dict):
    """
    Helper-functie om door paragrafen te lussen en tekst te vervangen.
    (Deze functie is ongewijzigd, maar roept de nieuwe 'slimme' helper aan)
    """
    for para in paragrafen:
        # Sla paragrafen over die geen placeholders lijken te hebben
        if '{' not in para.text:
            continue
            
        # Itereer over alle termen
        for term, input_waarde in vervangingen.items():
            # Check of de term in de *volledige* tekst van de paragraaf staat
            if term in para.text:
                # Zo ja, roep de slimme vervanger aan.
                vervang_tekst_met_behoud_van_opmaak(para, term, input_waarde)


def vervang_in_word(template_pad: str, output_pad: str, vervangingen: dict):
    """
    Opent het Word-template, voert alle vervangingen door en slaat het op als een nieuw bestand.
    (Deze functie blijft ongewijzigd)
    """
    try:
        doc = Document(template_pad)
        
        # 1. Vervang in standaardparagrafen
        verwerk_paragrafen(doc.paragraphs, vervangingen)
        
        # 2. Vervang in tabellen
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    verwerk_paragrafen(cell.paragraphs, vervangingen)
                        
        # 3. Vervang in kop- en voetteksten (headers/footers)
        for section in doc.sections:
            verwerk_paragrafen(section.header.paragraphs, vervangingen)
            verwerk_paragrafen(section.footer.paragraphs, vervangingen)
            verwerk_paragrafen(section.first_page_header.paragraphs, vervangingen)
            verwerk_paragrafen(section.first_page_footer.paragraphs, vervangingen)
        
        doc.save(output_pad)
        print(f"🎉 Document succesvol opgeslagen als: {output_pad}")
        
    except FileNotFoundError:
        print(f"❌ FOUT: Word-template niet gevonden op: {template_pad}")
    except PermissionError:
        print(f"❌ FOUT: Toegang geweigerd. Is het output-bestand ('{output_pad}') misschien nog geopend?")
    except Exception as e:
        print(f"❌ FOUT bij het verwerken van Word: {e}")

def genereer_output_pad(template_pad: str) -> str:
    """
    Genereert een output-pad in dezelfde map als het template,
    met '_ingevuld' toegevoegd aan de bestandsnaam.
    """
    map_pad = os.path.dirname(template_pad)
    bestandsnaam = os.path.basename(template_pad)
    naam_zonder_ext, _ = os.path.splitext(bestandsnaam)
    nieuwe_bestandsnaam = f"{naam_zonder_ext}_ingevuld.docx"
    return os.path.join(map_pad, nieuwe_bestandsnaam)


# --- HOOFDUITVOERING ---
# (Deze is ongewijzigd gebleven)
if __name__ == "__main__":
    print("--- Start Word Document Vuller ---")
    
    # Stap 1: Vraag de gebruiker om BEIDE bestanden
    print("Sleep het Word-template hierheen en druk op Enter:")
    template_pad = input("-> ").strip().strip("& ").strip("'\"")
    
    print("\nSleep nu het Excel-bestand hierheen en druk op Enter:")
    excel_pad = input("-> ").strip().strip("& ").strip("'\"")
    
    print("\nBezig met verwerken...")

    # Stap 2: Genereer het output-pad (nog steeds automatisch)
    output_pad = genereer_output_pad(template_pad)
    
    print(f"ℹ️  Word-template: {template_pad}")
    print(f"ℹ️  Excel-bron:    {excel_pad}")
    print(f"ℹ️  Output-bestand: {output_pad}\n")

    # Stap 3: Controleer of BEIDE invoerbestanden bestaan
    bestanden_gevonden = True
    if not os.path.exists(template_pad):
        print(f"❌ FOUT: Kan Word-template niet vinden op: {template_pad}")
        bestanden_gevonden = False
        
    if not os.path.exists(excel_pad):
        print(f"❌ FOUT: Kan Excel-bestand niet vinden op: {excel_pad}")
        bestanden_gevonden = False

    if bestanden_gevonden:
        # Stap 4: Laad de vervangingen uit Excel
        vervangingen_dict = laad_vervangingen(excel_pad)
        
        # Stap 5: Als het laden gelukt is, verwerk het Word-document
        if vervangingen_dict:
            vervang_in_word(template_pad, output_pad, vervangingen_dict)
            
    input("\nDruk op Enter om het venster te sluiten...")