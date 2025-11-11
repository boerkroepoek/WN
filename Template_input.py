import docx
import openpyxl
import re
from pathlib import Path  # Gebruik pathlib voor modern padbeheer
from openpyxl.styles import Font

def haal_tekst_uit_word(docx_pad):
    """
    Haalt alle tekst uit een .docx-bestand, inclusief paragrafen en tabellen.
    """
    try:
        doc = docx.Document(docx_pad)
        volledige_tekst = []

        # Haal tekst uit paragrafen
        for para in doc.paragraphs:
            volledige_tekst.append(para.text)

        # Haal tekst uit tabellen
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        volledige_tekst.append(para.text)
                        
        return "\n".join(volledige_tekst)

    except Exception as e:
        print(f"Fout bij het lezen van het Word-bestand: {e}")
        return None

def vind_woorden_tussen_haken(tekst):
    """
    Vindt alle tekststrings die tussen { } staan met behulp van regex.
    """
    # AANPASSING: De (.*?) is vervangen door .*?
    # Dit zorgt ervoor dat de regex de VOLLEDIGE match pakt (incl. haken),
    # in plaats van alleen de 'capturing group' (de tekst ertussen).
    # re.DOTALL zorgt dat het ook werkt als de term over meerdere regels staat.
    patroon = re.compile(r'\{.*?\}', re.DOTALL)
    gevonden_woorden = re.findall(patroon, tekst)
    
    # We hoeven .strip() niet meer te doen, we willen de volledige match.
    return gevonden_woorden

def schrijf_naar_excel(termen_lijst, excel_pad):
    """
    Schrijft de gevonden unieke termen naar een Excel-bestand in één kolom.
    """
    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Gevonden Termen"

        # AANPASSING: Header-tekst verduidelijkt.
        ws["A1"] = "Gevonden Term"
        ws["A1"].font = Font(bold=True)
        ws["B1"] = "input"
        ws["B1"].font = Font(bold=True)

        for index, term in enumerate(termen_lijst, start=2):
            ws[f"A{index}"] = term

        ws.column_dimensions['A'].width = 50
        wb.save(excel_pad)
        print(f"Excel-bestand succesvol opgeslagen als: {excel_pad}")

    except PermissionError:
        print(f"\nFOUT: Kan niet schrijven naar {excel_pad}.")
        print("Het bestand is mogelijk al geopend in Excel. Sluit het en probeer opnieuw.")
    except Exception as e:
        print(f"Fout bij het schrijven naar het Excel-bestand: {e}")

def verwerk_document(word_invoer_pad):
    """
    Hoofdfunctie om het hele proces te doorlopen.
    """
    invoer_pad = Path(word_invoer_pad)

    print(f"1. Bezig met lezen van: {invoer_pad.name}")
    document_tekst = haal_tekst_uit_word(invoer_pad)

    if document_tekst:
        print("2. Tekst geëxtraheerd. Bezig met zoeken naar termen (inclusief {})...")
        gevonden_termen = vind_woorden_tussen_haken(document_tekst)
        
        if not gevonden_termen:
            print("Geen termen tussen {} gevonden.")
            return

        print(f"   ... {len(gevonden_termen)} instanties gevonden.")

        unieke_termen = list(dict.fromkeys(gevonden_termen))
        print(f"   ... {len(unieke_termen)} unieke termen gevonden.")

        basis_naam = invoer_pad.stem
        excel_uitvoer = invoer_pad.parent / f"{basis_naam}_termen.xlsx"

        print(f"3. Bezig met schrijven naar: {excel_uitvoer}")
        schrijf_naar_excel(unieke_termen, excel_uitvoer)
        
        print("\nVerwerking voltooid!")

# --- HOE TE GEBRUIKEN ---
if __name__ == "__main__":
    pad_input = input("Voer het volledige pad naar het Word (.docx) bestand in: ")
    pad_input = pad_input.strip().strip('"').strip("'")
    
    word_pad = Path(pad_input)

    if not word_pad.exists():
        print(f"Fout: Kan het bestand niet vinden op het opgegeven pad:")
        print(f"{word_pad}")
        print("Controleer het pad en probeer het opnieuw.")
    elif word_pad.suffix.lower() != '.docx':
        print(f"Fout: Het opgegeven bestand is geen .docx bestand (eindigt niet op .docx).")
    else:
        verwerk_document(word_pad)